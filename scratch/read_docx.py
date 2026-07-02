import docx
import os

files = [
    r"c:\Users\Nodar\Batumi Hills\01-Docs\ბათუმი_ჰილსის_განშლის_შედარება.docx",
    r"c:\Users\Nodar\Batumi Hills\01-Docs\ბათუმი_ჰილსის_ძველი_განშლა_კორექტირებადი.docx",
    r"c:\Users\Nodar\2026 antigraviti\გალთ_გრუპის_ვალდებულებები_შესწორებული.docx"
]

output_file = r"c:\Users\Nodar\2026 antigraviti\scratch\docx_text_dump.txt"

with open(output_file, "w", encoding="utf-8") as out:
    for f in files:
        if os.path.exists(f):
            out.write(f"\n=================== {os.path.basename(f)} ===================\n")
            try:
                doc = docx.Document(f)
                out.write(f"Paragraphs: {len(doc.paragraphs)}\n")
                for p in doc.paragraphs:
                    if p.text.strip():
                        out.write(p.text + "\n")
                
                out.write("\n--- Tables ---\n")
                for i, table in enumerate(doc.tables):
                    out.write(f"\nTable {i+1}:\n")
                    for row in table.rows:
                        row_text = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                        out.write(" | ".join(row_text) + "\n")
            except Exception as e:
                out.write(f"Error: {e}\n")
        else:
            out.write(f"Path not found: {f}\n")

print("Done! Written to docx_text_dump.txt")
